#!/usr/bin/env python3
"""
build.py: Compile LaTeX to PDF and/or build Word (.docx) document.

Usage:
    python build.py --input paper.tex --format pdf
    python build.py --input paper.tex --format docx
    python build.py --input paper.tex --format both

Requirements:
    - pdflatex (from MacTeX, TinyTeX, or TeX Live) for PDF output
    - bibtex for bibliography processing
    - python-docx (pip install python-docx) for Word output

NOTE: Word output is built natively with python-docx, NOT pandoc.
Pandoc corrupts .docx files when converting tikz diagrams and complex
LaTeX math environments. python-docx builds proper Office XML directly.
"""

import argparse
import subprocess
import shutil
import sys
import os
import re
import json


def check_dependency(name):
    """Check if a command-line tool is available."""
    if shutil.which(name) is None:
        print(f"ERROR: '{name}' is not installed or not in PATH.")
        if name == "pdflatex":
            print("  Install via: brew install --cask mactex")
            print("  Or lightweight: brew install --cask basictex")
        elif name == "bibtex":
            print("  Usually included with your LaTeX distribution.")
        return False
    return True


def check_python_docx():
    """Check if python-docx is available."""
    try:
        import docx
        return True
    except ImportError:
        print("ERROR: python-docx is not installed.")
        print("  Install via: pip install python-docx")
        return False


def compile_pdf(tex_path):
    """Compile .tex to .pdf using pdflatex + bibtex."""
    if not check_dependency("pdflatex") or not check_dependency("bibtex"):
        return False

    work_dir = os.path.dirname(os.path.abspath(tex_path))
    base_name = os.path.splitext(os.path.basename(tex_path))[0]

    commands = [
        ["pdflatex", "-interaction=nonstopmode", "-output-directory", work_dir, tex_path],
        ["bibtex", os.path.join(work_dir, base_name)],
        ["pdflatex", "-interaction=nonstopmode", "-output-directory", work_dir, tex_path],
        ["pdflatex", "-interaction=nonstopmode", "-output-directory", work_dir, tex_path],
    ]

    for cmd in commands:
        try:
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                cwd=work_dir,
                timeout=60
            )
            if result.returncode != 0 and "bibtex" not in cmd[0]:
                print(f"Warning: {cmd[0]} returned code {result.returncode}")
                if result.stderr:
                    print(f"  stderr: {result.stderr[:500]}")
        except FileNotFoundError:
            print(f"ERROR: Command not found: {cmd[0]}")
            return False
        except subprocess.TimeoutExpired:
            print(f"ERROR: {cmd[0]} timed out after 60 seconds.")
            return False

    pdf_path = os.path.join(work_dir, f"{base_name}.pdf")
    if os.path.exists(pdf_path):
        print(f"PDF created: {pdf_path}")
        return True
    else:
        print("ERROR: PDF was not created. Check LaTeX log for errors.")
        log_path = os.path.join(work_dir, f"{base_name}.log")
        if os.path.exists(log_path):
            print(f"  Log file: {log_path}")
        return False


def build_docx_from_json(json_path, docx_path):
    """
    Build a Word document natively using python-docx.

    Expects a JSON file with the document structure:
    {
        "title": "Paper Title",
        "author": "Author Name",
        "date": "2026-04-09",
        "sections": [
            {
                "heading": "Section Title",
                "level": 1,
                "content": [
                    {"type": "paragraph", "text": "Paragraph text..."},
                    {"type": "equation", "latex": "Y_{it} = ...", "label": "Eq. (1)"},
                    {"type": "table", "headers": [...], "rows": [...]},
                    {"type": "flowchart_description", "steps": [
                        {"name": "Step 1", "detail": "...", "color": "blue"}
                    ]},
                    {"type": "checklist", "items": [
                        {"text": "Check item", "checked": false}
                    ]}
                ]
            }
        ],
        "references": [
            {"key": "smith2020", "text": "Smith, J. (2020). Title. Journal."}
        ]
    }
    """
    if not check_python_docx():
        return False

    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    try:
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError) as e:
        print(f"ERROR: Could not read JSON file: {e}")
        return False

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title = data.get("title", "Untitled")
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Author and date
    author = data.get("author", "")
    date_str = data.get("date", "")
    if author or date_str:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if author:
            run = meta_para.add_run(author)
            run.font.size = Pt(12)
        if author and date_str:
            meta_para.add_run("\n")
        if date_str:
            run = meta_para.add_run(date_str)
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(100, 100, 100)

    # Process sections
    for section in data.get("sections", []):
        level = section.get("level", 1)
        heading = section.get("heading", "")
        if heading:
            doc.add_heading(heading, level=min(level, 4))

        for content in section.get("content", []):
            content_type = content.get("type", "paragraph")

            if content_type == "paragraph":
                text = content.get("text", "")
                para = doc.add_paragraph(text)
                para.style.font.size = Pt(12)

            elif content_type == "equation":
                # Render equation in a styled paragraph
                latex = content.get("latex", "")
                label = content.get("label", "")
                eq_para = doc.add_paragraph()
                eq_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Equation in Cambria Math font
                run = eq_para.add_run(latex)
                run.font.name = "Cambria Math"
                run.font.size = Pt(11)
                if label:
                    run = eq_para.add_run(f"    {label}")
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(100, 100, 100)

            elif content_type == "table":
                headers = content.get("headers", [])
                rows = content.get("rows", [])
                if headers:
                    table = doc.add_table(
                        rows=1 + len(rows),
                        cols=len(headers)
                    )
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    # Header row
                    for i, header in enumerate(headers):
                        cell = table.rows[0].cells[i]
                        cell.text = header
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(10)
                    # Data rows
                    for r_idx, row in enumerate(rows):
                        for c_idx, cell_text in enumerate(row):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = str(cell_text)
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.size = Pt(10)
                    doc.add_paragraph()  # spacing after table

            elif content_type == "flowchart_description":
                # Render methodology pipeline as a colour-coded table
                steps = content.get("steps", [])
                color_map = {
                    "green": "C8E6C8",
                    "blue": "C8D7F0",
                    "orange": "FFE1C8",
                    "purple": "E1D2F0",
                    "gray": "E6E6E6",
                }
                if steps:
                    table = doc.add_table(rows=len(steps), cols=3)
                    table.style = "Table Grid"
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, step in enumerate(steps):
                        # Step number
                        cell_num = table.rows[i].cells[0]
                        cell_num.text = str(i + 1)
                        # Step name
                        cell_name = table.rows[i].cells[1]
                        cell_name.text = step.get("name", "")
                        for p in cell_name.paragraphs:
                            for r in p.runs:
                                r.bold = True
                                r.font.size = Pt(10)
                        # Step detail
                        cell_detail = table.rows[i].cells[2]
                        cell_detail.text = step.get("detail", "")
                        for p in cell_detail.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(10)
                        # Apply colour
                        color_name = step.get("color", "blue")
                        bg_color = color_map.get(color_name, color_map["blue"])
                        from docx.oxml.ns import qn
                        hex_color = color_map.get(color_name, "C8D7F0")
                        for cell in [cell_num, cell_name, cell_detail]:
                            tc_pr = cell._element.get_or_add_tcPr()
                            shading_elem = tc_pr.makeelement(
                                qn("w:shd"),
                                {
                                    qn("w:fill"): hex_color,
                                    qn("w:val"): "clear",
                                }
                            )
                            tc_pr.append(shading_elem)
                    doc.add_paragraph()  # spacing after flowchart

            elif content_type == "checklist":
                items = content.get("items", [])
                for item in items:
                    checked = item.get("checked", False)
                    marker = "[x]" if checked else "[ ]"
                    text = item.get("text", "")
                    para = doc.add_paragraph(f"{marker} {text}")
                    para.style.font.size = Pt(11)

    # References section
    references = data.get("references", [])
    if references:
        doc.add_heading("References", level=1)
        for ref in references:
            text = ref.get("text", "")
            para = doc.add_paragraph(text)
            para.style.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(4)

    # Save
    try:
        doc.save(docx_path)
        print(f"Word document created: {docx_path}")
        return True
    except Exception as e:
        print(f"ERROR: Could not save Word document: {e}")
        return False


def build_docx_from_tex(tex_path):
    """
    Build Word document from .tex file.

    Strategy:
    1. Look for a companion .json file (same name, .json extension).
       If found, use build_docx_from_json() for native Word output.
    2. If no .json file, fall back to pandoc (for simple LaTeX without
       tikz or complex math).
    """
    work_dir = os.path.dirname(os.path.abspath(tex_path))
    base_name = os.path.splitext(os.path.basename(tex_path))[0]
    json_path = os.path.join(work_dir, f"{base_name}.json")
    docx_path = os.path.join(work_dir, f"{base_name}.docx")

    # Prefer JSON-based native Word build
    if os.path.exists(json_path):
        print("Found companion .json file. Building Word natively with python-docx.")
        return build_docx_from_json(json_path, docx_path)

    # Check if the .tex has tikz or complex environments
    try:
        with open(tex_path, "r", encoding="utf-8") as f:
            tex_content = f.read()
    except Exception:
        tex_content = ""

    has_tikz = "\\begin{tikzpicture}" in tex_content
    has_complex = has_tikz or "\\usetikzlibrary" in tex_content

    if has_complex:
        print("WARNING: LaTeX contains tikz diagrams or complex environments.")
        print("  pandoc will produce a corrupted .docx for these elements.")
        print("  To get a proper Word file, generate a companion .json file")
        print("  with the document structure. See build.py documentation.")
        print("  Attempting pandoc conversion anyway (diagrams will be missing)...")

    # Fallback to pandoc for simple documents
    if not shutil.which("pandoc"):
        print("ERROR: pandoc is not installed and no .json file found.")
        print("  Install pandoc: brew install pandoc")
        print("  Or generate a .json file for native Word output.")
        return False

    bib_path = os.path.join(work_dir, "references.bib")
    cmd = [
        "pandoc", tex_path,
        "-o", docx_path,
        "--from", "latex",
        "--to", "docx",
    ]
    if os.path.exists(bib_path):
        cmd.extend(["--citeproc", "--bibliography", bib_path])

    try:
        result = subprocess.run(
            cmd, capture_output=True, text=True,
            cwd=work_dir, timeout=60
        )
        if result.returncode != 0:
            print(f"ERROR: pandoc failed with code {result.returncode}")
            if result.stderr:
                print(f"  stderr: {result.stderr[:500]}")
            return False
    except FileNotFoundError:
        print("ERROR: pandoc not found.")
        return False
    except subprocess.TimeoutExpired:
        print("ERROR: pandoc timed out after 60 seconds.")
        return False

    if os.path.exists(docx_path):
        print(f"Word document created: {docx_path}")
        if has_complex:
            print("  NOTE: tikz diagrams were not converted. Use the PDF version")
            print("  for the complete document with diagrams.")
        return True
    else:
        print("ERROR: Word document was not created.")
        return False


def clean_aux_files(tex_path):
    """Remove auxiliary files created during compilation."""
    work_dir = os.path.dirname(os.path.abspath(tex_path))
    base_name = os.path.splitext(os.path.basename(tex_path))[0]
    extensions = [".aux", ".log", ".bbl", ".blg", ".out", ".toc", ".lof", ".lot"]

    for ext in extensions:
        aux_file = os.path.join(work_dir, f"{base_name}{ext}")
        if os.path.exists(aux_file):
            os.remove(aux_file)


def main():
    parser = argparse.ArgumentParser(
        description="Compile LaTeX to PDF and/or build Word document."
    )
    parser.add_argument(
        "--input", required=True,
        help="Path to the .tex file"
    )
    parser.add_argument(
        "--format", required=True, choices=["pdf", "docx", "both"],
        help="Output format: pdf, docx, or both"
    )
    parser.add_argument(
        "--clean", action="store_true",
        help="Remove auxiliary files after compilation"
    )

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: Input file not found: {args.input}")
        sys.exit(1)

    success = True

    if args.format in ("pdf", "both"):
        if not compile_pdf(args.input):
            success = False

    if args.format in ("docx", "both"):
        if not build_docx_from_tex(args.input):
            success = False

    if args.clean:
        clean_aux_files(args.input)

    if success:
        print("\nBuild completed successfully.")
    else:
        print("\nBuild completed with errors. Check output above.")
        sys.exit(1)


if __name__ == "__main__":
    main()
